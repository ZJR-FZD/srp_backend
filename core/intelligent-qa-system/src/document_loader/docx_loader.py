"""
Word 文档加载器
"""
from typing import List
from pathlib import Path
from docx import Document as DocxDocument

from .base_loader import BaseLoader, Document


class DOCXLoader(BaseLoader):
    """Word 文档加载器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
    
    def load(self, file_path: str) -> List[Document]:
        """
        加载 Word 文档
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            List[Document]: 文档列表（整个文档作为一个 Document）
        """
        if not self.is_supported(file_path):
            raise ValueError(f"不支持的文件类型: {file_path}")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        try:
            # 读取 Word 文档
            doc = DocxDocument(file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格内容（可选）
            tables_content = self._extract_tables(doc)
            
            # 合并内容
            content = "\n\n".join(paragraphs)
            if tables_content:
                content += "\n\n" + tables_content
            
            # 创建元数据
            metadata = self._create_metadata(
                file_path,
                paragraphs_count=len(paragraphs),
                tables_count=len(doc.tables)
            )
            
            # 添加文档属性（如果有）
            if doc.core_properties.title:
                metadata['title'] = doc.core_properties.title
            if doc.core_properties.author:
                metadata['author'] = doc.core_properties.author
            
            return [Document(content=content, metadata=metadata)]
        
        except Exception as e:
            raise Exception(f"Word 文档加载失败: {e}")
    
    def _extract_tables(self, doc: DocxDocument) -> str:
        """提取表格内容"""
        if not doc.tables:
            return ""
        
        tables_text = []
        
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # 只添加非空行
                    table_data.append(" | ".join(row_data))
            
            if table_data:
                tables_text.append(f"表格 {i+1}:\n" + "\n".join(table_data))
        
        return "\n\n".join(tables_text)
    
    def load_by_sections(self, file_path: str) -> List[Document]:
        """
        按段落加载 Word 文档（每个段落作为一个 Document）
        适用于需要更细粒度切分的场景
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            List[Document]: 文档列表（每个段落一个 Document）
        """
        if not self.is_supported(file_path):
            raise ValueError(f"不支持的文件类型: {file_path}")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        documents = []
        
        try:
            doc = DocxDocument(file_path)
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    metadata = self._create_metadata(
                        file_path,
                        paragraph_index=i + 1,
                        total_paragraphs=len(doc.paragraphs)
                    )
                    
                    documents.append(Document(
                        content=text,
                        metadata=metadata
                    ))
        
        except Exception as e:
            raise Exception(f"Word 文档加载失败: {e}")
        
        return documents